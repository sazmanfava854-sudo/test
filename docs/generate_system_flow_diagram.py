#!/usr/bin/env python3
"""Generate system flow diagram (SVG, PNG, DOCX) for Word import."""

import base64
from pathlib import Path

import arabic_reshaper
import matplotlib.pyplot as plt
import matplotlib.patches as mpatches
from bidi.algorithm import get_display
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt
from matplotlib import font_manager
from matplotlib.path import Path as MplPath

DOCS_DIR = Path(__file__).resolve().parent
FONTS_DIR = DOCS_DIR / "fonts"
SVG_PATH = DOCS_DIR / "system_flow_diagram.svg"
PNG_PATH = DOCS_DIR / "system_flow_diagram.png"
DOCX_PATH = DOCS_DIR / "system_flow_diagram.docx"
FONT_PATH = "/usr/share/fonts/truetype/noto/NotoNaskhArabic-Regular.ttf"
FONT_BOLD_PATH = "/usr/share/fonts/truetype/noto/NotoNaskhArabic-Bold.ttf"
NAZANIN_PATH = FONTS_DIR / "BNazanin.ttf"
DPI = 300
PERSIAN_DIGITS = "۰۱۲۳۴۵۶۷۸۹"

font_manager.fontManager.addfont(FONT_PATH)
font_manager.fontManager.addfont(FONT_BOLD_PATH)
font_manager.fontManager.addfont(str(NAZANIN_PATH))
FONT_REG = font_manager.FontProperties(family=["Noto Naskh Arabic", "DejaVu Sans"])
FONT_BOLD = font_manager.FontProperties(family=["Noto Naskh Arabic", "DejaVu Sans"], weight="bold")
FONT_NAZANIN = font_manager.FontProperties(fname=str(NAZANIN_PATH))
FONT_LATIN = font_manager.FontProperties(family="DejaVu Sans")
FONT_LATIN_BOLD = font_manager.FontProperties(family="DejaVu Sans", weight="bold")


def fa(text: str, for_matplotlib: bool = False) -> str:
    shaped = arabic_reshaper.reshape(text)
    return shaped if for_matplotlib else get_display(shaped)


def fap(text: str) -> str:
    return fa(text, for_matplotlib=True)


def persian_num(n: int) -> str:
    return "".join(PERSIAN_DIGITS[int(d)] for d in str(n))


def algo_parts(number: int, suffix: str = "") -> list[tuple[str, font_manager.FontProperties]]:
    parts: list[tuple[str, font_manager.FontProperties]] = [
        (fap("الگوریتم "), FONT_REG),
        (persian_num(number), FONT_NAZANIN),
    ]
    if suffix:
        ascii_ratio = sum(1 for c in suffix if ord(c) < 128) / max(len(suffix), 1)
        fp = FONT_LATIN if ascii_ratio > 0.5 else FONT_REG
        parts.append((fap(f" — {suffix}") if fp is FONT_REG else f" — {suffix}", fp))
    return parts


def _pick_font(line: str, bold: bool = False):
    ascii_ratio = sum(1 for c in line if ord(c) < 128) / max(len(line), 1)
    if ascii_ratio > 0.6:
        return FONT_LATIN_BOLD if bold else FONT_LATIN
    return FONT_BOLD if bold else FONT_REG


def _measure_parts(ax, fig, parts, fs):
    fig.canvas.draw()
    renderer = fig.canvas.get_renderer()
    widths = []
    for text, fp in parts:
        tmp = ax.text(0, 0, text, fontproperties=fp, fontsize=fs)
        bb = tmp.get_window_extent(renderer=renderer)
        widths.append(bb.width)
        tmp.remove()
    return widths


def draw_mixed_line(ax, fig, cx, cy, parts, fs):
    widths = _measure_parts(ax, fig, parts, fs)
    total = sum(widths)
    x = cx - total / 2
    for (text, fp), w in zip(parts, widths):
        ax.text(x + w / 2, cy, text, ha="center", va="center", fontsize=fs, fontproperties=fp, color="#1a1a1a")
        x += w


def draw_rounded_box(ax, fig, cx, cy, w, h, fc, ec, lines, fs=11, bold=False, radius=0.08):
    box = mpatches.FancyBboxPatch(
        (cx - w / 2, cy - h / 2),
        w,
        h,
        boxstyle=f"round,pad=0.02,rounding_size={radius}",
        facecolor=fc,
        edgecolor=ec,
        linewidth=1.8,
    )
    ax.add_patch(box)
    line_h = 0.22 if len(lines) > 1 else 0
    start_y = cy + line_h * (len(lines) - 1) / 2
    for i, line in enumerate(lines):
        y = start_y - i * 0.28
        if isinstance(line, list):
            draw_mixed_line(ax, fig, cx, y, line, fs)
        else:
            ax.text(
                cx,
                y,
                line,
                ha="center",
                va="center",
                fontsize=fs,
                fontproperties=_pick_font(line, bold),
                color="#1a1a1a",
            )


def draw_diamond(ax, fig, cx, cy, size, fc, ec, lines, fs=10.5):
    half = size / 2
    verts = [(cx, cy + half), (cx + half, cy), (cx, cy - half), (cx - half, cy), (cx, cy + half)]
    codes = [MplPath.MOVETO, MplPath.LINETO, MplPath.LINETO, MplPath.LINETO, MplPath.CLOSEPOLY]
    patch = mpatches.PathPatch(MplPath(verts, codes), facecolor=fc, edgecolor=ec, linewidth=1.8)
    ax.add_patch(patch)
    start_y = cy + 0.12 * (len(lines) - 1)
    for i, line in enumerate(lines):
        y = start_y - i * 0.24
        if isinstance(line, list):
            draw_mixed_line(ax, fig, cx, y, line, fs)
        else:
            ax.text(
                cx,
                y,
                line,
                ha="center",
                va="center",
                fontsize=fs,
                fontproperties=_pick_font(line),
                color="#1a1a1a",
            )


def draw_arrow(ax, x1, y1, x2, y2):
    ax.annotate(
        "",
        xy=(x2, y2),
        xytext=(x1, y1),
        arrowprops=dict(arrowstyle="-|>", color="#455A64", lw=1.8, mutation_scale=14),
    )


def build_matplotlib_png(png_path: Path) -> None:
    fig, ax = plt.subplots(figsize=(7.2, 14.8), dpi=DPI)
    ax.set_xlim(0, 10)
    ax.set_ylim(0, 20.5)
    ax.axis("off")
    fig.patch.set_facecolor("#FAFAFA")

    cx = 5.0
    bw = 5.8

    ax.text(
        cx,
        20.0,
        fap("نمودار جریان کلی سیستم توصیه‌گر فناوری IoT"),
        ha="center",
        va="center",
        fontsize=14,
        fontproperties=FONT_BOLD,
        color="#263238",
    )

    y = 18.8
    draw_rounded_box(
        ax, fig, cx, y, bw, 0.95, "#E3F2FD", "#2196F3",
        [fap("داده خام فناوری‌ها"), fap("(۱۳ فناوری × ۸ معیار)")], fs=11.5, bold=True,
    )

    y -= 1.35
    draw_arrow(ax, cx, y + 0.62, cx, y + 0.18)
    draw_rounded_box(
        ax, fig, cx, y, bw, 1.05, "#E8EAF6", "#3F51B5",
        [algo_parts(1, "پیش‌پردازش"), "Log1p + Z-Score"], fs=11.2, bold=True,
    )

    y -= 1.45
    draw_arrow(ax, cx, y + 0.67, cx, y + 0.2)
    draw_rounded_box(
        ax, fig, cx, y, bw, 1.05, "#E8EAF6", "#3F51B5",
        [algo_parts(2, "K-Means"), fap("انتخاب خودکار k + Silhouette")], fs=11.2, bold=True,
    )

    y -= 1.55
    draw_arrow(ax, cx, y + 0.72, cx, y + 0.45)
    draw_diamond(ax, fig, cx, y, 1.55, "#FFF3E0", "#FF9800", [fap("انتخاب خوشه"), fap("توسط کاربر")])

    y -= 1.55
    draw_arrow(ax, cx, y + 0.72, cx, y + 0.2)
    draw_rounded_box(
        ax, fig, cx, y, bw, 0.95, "#F3E5F5", "#9C27B0",
        [algo_parts(3, "AHP"), fap("وزن پایه w")], fs=11.2, bold=True,
    )

    y -= 1.35
    draw_arrow(ax, cx, y + 0.62, cx, y + 0.18)
    draw_rounded_box(
        ax, fig, cx, y, bw, 0.82, "#FCE4EC", "#E91E63",
        [fap("پرسشنامه سناریو")], fs=11.5, bold=True,
    )

    y -= 1.25
    draw_arrow(ax, cx, y + 0.56, cx, y + 0.18)
    draw_rounded_box(
        ax, fig, cx, y, bw, 0.95, "#F3E5F5", "#9C27B0",
        [algo_parts(4, "تعدیل نمایی"), fap("وزن w̃")], fs=11.2, bold=True,
    )

    y_rank = y - 1.55
    sw = 2.55
    sx = [2.15, 5.0, 7.85]
    for i, (x, label) in enumerate(zip(sx, ["TOPSIS", "VIKOR", "COPRAS"])):
        draw_arrow(ax, cx, y - 0.55, x, y_rank + 0.55)
        draw_rounded_box(
            ax, fig, x, y_rank, sw, 1.0, "#E0F2F1", "#009688",
            [algo_parts(i + 5), label], fs=10.5, bold=True,
        )

    y_spear = y_rank - 1.45
    for x in sx:
        draw_arrow(ax, x, y_rank - 0.58, cx, y_spear + 0.55)
    draw_rounded_box(
        ax, fig, cx, y_spear, bw, 1.15, "#FFF8E1", "#FFC107",
        [
            algo_parts(8, "Spearman"),
            fap("سه مقایسه زوجی:"),
            "TOPSIS↔VIKOR | TOPSIS↔COPRAS | VIKOR↔COPRAS",
        ],
        fs=10.2,
        bold=True,
    )

    y_final = y_spear - 1.45
    draw_arrow(ax, cx, y_spear - 0.65, cx, y_final + 0.35)
    draw_rounded_box(
        ax, fig, cx, y_final, 3.8, 0.9, "#C8E6C9", "#4CAF50",
        [fap("توصیه نهایی")], fs=14, bold=True, radius=0.25,
    )

    fig.savefig(png_path, dpi=DPI, bbox_inches="tight", facecolor=fig.get_facecolor(), pad_inches=0.15)
    plt.close(fig)


def _svg_algo_line(number: int, suffix: str = "") -> str:
    nazanin = persian_num(number)
    prefix = fa("الگوریتم ")
    if suffix:
        ascii_ratio = sum(1 for c in suffix if ord(c) < 128) / max(len(suffix), 1)
        if ascii_ratio > 0.5:
            tail = f" — {suffix}"
        else:
            tail = fa(f" — {suffix}")
        return (
            f'<tspan font-family="NotoNaskhArabic">{prefix}</tspan>'
            f'<tspan font-family="BNazanin">{nazanin}</tspan>'
            f'<tspan font-family="{"DejaVu Sans" if ascii_ratio > 0.5 else "NotoNaskhArabic"}">{tail}</tspan>'
        )
    return (
        f'<tspan font-family="NotoNaskhArabic">{prefix}</tspan>'
        f'<tspan font-family="BNazanin">{nazanin}</tspan>'
    )


def build_svg() -> str:
    font_data = base64.b64encode(Path(FONT_PATH).read_bytes()).decode("ascii")
    font_bold_data = base64.b64encode(Path(FONT_BOLD_PATH).read_bytes()).decode("ascii")
    nazanin_data = base64.b64encode(NAZANIN_PATH.read_bytes()).decode("ascii")
    w, h = 720, 1480

    def box(x, y, bw, bh, fill, stroke, lines, fs=15, bold=False, radius=10, mixed_first=False):
        fw = "bold" if bold else "normal"
        text_y = y + bh / 2 - (len(lines) - 1) * 9
        texts = ""
        for i, line in enumerate(lines):
            if isinstance(line, tuple) and line[0] == "algo":
                inner = _svg_algo_line(line[1], line[2] if len(line) > 2 else "")
                texts += (
                    f'<text x="{x + bw/2}" y="{text_y + i * 22}" text-anchor="middle" '
                    f'font-size="{fs}" font-weight="{fw}" fill="#1a1a1a">{inner}</text>\n'
                )
            else:
                texts += (
                    f'<text x="{x + bw/2}" y="{text_y + i * 22}" text-anchor="middle" '
                    f'font-family="NotoNaskhArabic" font-size="{fs}" font-weight="{fw}" '
                    f'fill="#1a1a1a">{line}</text>\n'
                )
        return (
            f'<rect x="{x}" y="{y}" width="{bw}" height="{bh}" rx="{radius}" ry="{radius}" '
            f'fill="{fill}" stroke="{stroke}" stroke-width="2"/>\n{texts}'
        )

    def diamond(cx, cy, size, fill, stroke, lines, fs=14):
        half = size / 2
        pts = f"{cx},{cy-half} {cx+half},{cy} {cx},{cy+half} {cx-half},{cy}"
        texts = ""
        text_y = cy - (len(lines) - 1) * 9
        for i, line in enumerate(lines):
            texts += (
                f'<text x="{cx}" y="{text_y + i * 20}" text-anchor="middle" '
                f'font-family="NotoNaskhArabic" font-size="{fs}" fill="#1a1a1a">{line}</text>\n'
            )
        return f'<polygon points="{pts}" fill="{fill}" stroke="{stroke}" stroke-width="2"/>\n{texts}'

    def arrow(x1, y1, x2, y2):
        return (
            f'<line x1="{x1}" y1="{y1}" x2="{x2}" y2="{y2}" stroke="#455A64" '
            f'stroke-width="2.2" marker-end="url(#arrow)"/>'
        )

    cx = w // 2
    bw = 420
    bx = (w - bw) // 2
    sbw, gap = 190, 25
    sx1 = (w - 3 * sbw - 2 * gap) // 2
    sx2 = sx1 + sbw + gap
    sx3 = sx2 + sbw + gap

    elements = [
        box(bx, 40, bw, 58, "#E3F2FD", "#2196F3",
            [fa("داده خام فناوری‌ها"), fa("(۱۳ فناوری × ۸ معیار)")], fs=16, bold=True),
        arrow(cx, 98, cx, 118),
        box(bx, 118, bw, 72, "#E8EAF6", "#3F51B5",
            [("algo", 1, "پیش‌پردازش"), fa("Log1p + Z-Score")], fs=15, bold=True),
        arrow(cx, 190, cx, 220),
        box(bx, 220, bw, 72, "#E8EAF6", "#3F51B5",
            [("algo", 2, "K-Means"), fa("انتخاب خودکار k + Silhouette")], fs=15, bold=True),
        arrow(cx, 292, cx, 330),
        diamond(cx, 370, 130, "#FFF3E0", "#FF9800", [fa("انتخاب خوشه"), fa("توسط کاربر")]),
        arrow(cx, 435, cx, 465),
        box(bx, 465, bw, 58, "#F3E5F5", "#9C27B0",
            [("algo", 3, "AHP"), fa("وزن پایه w")], fs=15, bold=True),
        arrow(cx, 523, cx, 553),
        box(bx, 553, bw, 52, "#FCE4EC", "#E91E63", [fa("پرسشنامه سناریو")], fs=15, bold=True),
        arrow(cx, 605, cx, 635),
        box(bx, 635, bw, 58, "#F3E5F5", "#9C27B0",
            [("algo", 4, "تعدیل نمایی"), fa("وزن w̃")], fs=15, bold=True),
        arrow(cx, 693, cx, 730),
        box(sx1, 740, sbw, 68, "#E0F2F1", "#009688", [("algo", 5), "TOPSIS"], fs=14, bold=True),
        box(sx2, 740, sbw, 68, "#E0F2F1", "#009688", [("algo", 6), "VIKOR"], fs=14, bold=True),
        box(sx3, 740, sbw, 68, "#E0F2F1", "#009688", [("algo", 7), "COPRAS"], fs=14, bold=True),
        arrow(cx, 693, sx1 + sbw // 2, 730),
        arrow(cx, 693, sx2 + sbw // 2, 730),
        arrow(cx, 693, sx3 + sbw // 2, 730),
        arrow(sx1 + sbw // 2, 808, cx, 840),
        arrow(sx2 + sbw // 2, 808, cx, 840),
        arrow(sx3 + sbw // 2, 808, cx, 840),
        arrow(cx, 840, cx, 860),
        box(bx, 860, bw, 78, "#FFF8E1", "#FFC107",
            [("algo", 8, "Spearman"), fa("سه مقایسه زوجی:"),
             "TOPSIS↔VIKOR | TOPSIS↔COPRAS | VIKOR↔COPRAS"], fs=13, bold=True),
        arrow(cx, 938, cx, 978),
        box(bx + 60, 978, bw - 120, 58, "#C8E6C9", "#4CAF50", [fa("توصیه نهایی")], fs=18, bold=True, radius=28),
    ]

    body = "\n".join(elements)
    return f"""<?xml version="1.0" encoding="UTF-8"?>
<svg xmlns="http://www.w3.org/2000/svg" width="{w}" height="{h}" viewBox="0 0 {w} {h}">
  <defs>
    <marker id="arrow" markerWidth="10" markerHeight="10" refX="8" refY="3" orient="auto">
      <polygon points="0 0, 10 3, 0 6" fill="#455A64"/>
    </marker>
    <style>
      @font-face {{
        font-family: 'NotoNaskhArabic';
        src: url(data:font/truetype;base64,{font_data}) format('truetype');
        font-weight: normal;
      }}
      @font-face {{
        font-family: 'NotoNaskhArabic';
        src: url(data:font/truetype;base64,{font_bold_data}) format('truetype');
        font-weight: bold;
      }}
      @font-face {{
        font-family: 'BNazanin';
        src: url(data:font/truetype;base64,{nazanin_data}) format('truetype');
        font-weight: normal;
      }}
    </style>
  </defs>
  <rect width="100%" height="100%" fill="#FAFAFA"/>
  <text x="{cx}" y="28" text-anchor="middle" font-family="NotoNaskhArabic" font-size="20"
        font-weight="bold" fill="#263238">{fa("نمودار جریان کلی سیستم توصیه‌گر فناوری IoT")}</text>
  {body}
</svg>
"""


def build_docx(png_path: Path, docx_path: Path) -> None:
    doc = Document()
    for section in doc.sections:
        section.page_height = Cm(29.7)
        section.page_width = Cm(21.0)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2.5)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("نمودار جریان کلی سیستم")
    run.bold = True
    run.font.size = Pt(16)
    run.font.name = "B Nazanin"
    r = run._element.rPr
    if r is not None:
        r.rFonts.set(qn("w:cs"), "B Nazanin")

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run().add_picture(str(png_path), width=Inches(5.8))

    caption = doc.add_paragraph()
    caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap_run = caption.add_run("شکل — جریان داده از پیش‌پردازش تا توصیه نهایی")
    cap_run.font.size = Pt(12)
    cap_run.font.name = "B Nazanin"

    doc.save(docx_path)


def main() -> None:
    svg_content = build_svg()
    SVG_PATH.write_text(svg_content, encoding="utf-8")
    build_matplotlib_png(PNG_PATH)
    build_docx(PNG_PATH, DOCX_PATH)
    print(f"SVG:  {SVG_PATH}")
    print(f"PNG:  {PNG_PATH}")
    print(f"DOCX: {DOCX_PATH}")


if __name__ == "__main__":
    main()
