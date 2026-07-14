#!/usr/bin/env python3
"""Generate compact thesis questionnaire table (Word + TSV for copy-paste)."""

from pathlib import Path

from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt

OUTPUT_DOCX = Path(__file__).resolve().parent / "thesis_questions_table_compact.docx"
OUTPUT_TSV = Path(__file__).resolve().parent / "thesis_questions_table_compact.tsv"

FONT = "B Nazanin"

ROWS = [
    (
        "مساحت کل سایت استقرار چقدر است؟",
        "کوچک (کمتر از ۱ هکتار)",
        "برد انتقال",
        "−۱",
        "پوشش وسیع ارتباطی در اولویت طراحی شبکه نیست.",
    ),
    (
        "مساحت کل سایت استقرار چقدر است؟",
        "کوچک (کمتر از ۱ هکتار)",
        "بودجه پیوند",
        "−۱",
        "نیازی به توان پیوند بالا برای غلبه بر فواصل طولانی نیست.",
    ),
    (
        "مساحت کل سایت استقرار چقدر است؟",
        "بزرگ (بیش از ۱۰ هکتار)",
        "برد انتقال",
        "+۱",
        "پوشش گسترده‌تر محل استقرار، اهمیت برد انتقال بالاتر را افزایش می‌دهد.",
    ),
    (
        "مساحت کل سایت استقرار چقدر است؟",
        "بزرگ (بیش از ۱۰ هکتار)",
        "بودجه پیوند",
        "+۱",
        "فواصل بیشتر بین نقاط، نیاز به بودجه پیوند قوی‌تر را برجسته می‌کند.",
    ),
    (
        "شکل و گستره محیط چگونه است؟",
        "کوچک و متمرکز",
        "برد انتقال",
        "−۱",
        "در محیط متمرکز، برد زیاد اولویت اصلی انتخاب فناوری نیست.",
    ),
    (
        "منبع تغذیه گره‌ها چیست؟",
        "باتری محدود",
        "مصرف انرژی",
        "+۱",
        "فناوری‌های کم‌مصرف اهمیت بیشتری در انتخاب پیدا می‌کنند.",
    ),
    (
        "منبع تغذیه گره‌ها چیست؟",
        "برق شهری پایدار",
        "مصرف انرژی",
        "−۱",
        "با دسترسی پایدار به برق، محدودیت انرژی اهمیت کمتری می‌یابد.",
    ),
    (
        "محدودیت بودجه چگونه است؟",
        "بودجه سخت‌افزار محدود",
        "هزینه سرمایه‌ای سخت‌افزار",
        "+۱",
        "هزینه اولیه سخت‌افزار باید در وزن‌دهی اهمیت بیشتری بگیرد.",
    ),
    (
        "هزینه اتصال سالانه چقدر مهم است؟",
        "بسیار مهم",
        "هزینه عملیاتی سالانه اتصال",
        "+۱",
        "فناوری با هزینه اشتراک و اتصال پایین‌تر مطلوب‌تر است.",
    ),
    (
        "آیا ارتباط بلادرنگ لازم است؟",
        "بله",
        "تأخیر انتقال داده",
        "+۱",
        "تأخیر کم در انتقال داده اهمیت بیشتری در انتخاب فناوری دارد.",
    ),
    (
        "آیا پوشش اپراتوری قابل اتکا وجود دارد؟",
        "بله",
        "پشتیبانی شبکه سلولی",
        "−۱",
        "با پوشش مناسب اپراتور، وابستگی تصمیم به این معیار کاهش می‌یابد.",
    ),
    (
        "آیا داده حجیم ارسال می‌شود؟",
        "بله",
        "نرخ داده",
        "+۱",
        "فناوری‌های با ظرفیت انتقال بالاتر برای این سناریو مناسب‌ترند.",
    ),
]

HEADERS = ["پرسش", "پاسخ کاربر", "معیار تحت تأثیر", "اثر قاعده", "تفسیر"]


def set_rtl(paragraph, align=WD_ALIGN_PARAGRAPH.RIGHT):
    p_pr = paragraph._p.get_or_add_pPr()
    bidi = OxmlElement("w:bidi")
    bidi.set(qn("w:val"), "1")
    p_pr.append(bidi)
    paragraph.alignment = align


def set_cell_rtl(cell):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_pr.append(OxmlElement("w:bidiVisual"))
    v_align = OxmlElement("w:vAlign")
    v_align.set(qn("w:val"), "center")
    tc_pr.append(v_align)
    for p in cell.paragraphs:
        set_rtl(p)


def write_cell(cell, text, *, bold=False, center=False):
    set_cell_rtl(cell)
    p = cell.paragraphs[0]
    p.clear()
    set_rtl(p, WD_ALIGN_PARAGRAPH.CENTER if center else WD_ALIGN_PARAGRAPH.RIGHT)
    run = p.add_run(text)
    run.bold = bold
    run.font.name = FONT
    run.font.size = Pt(11)


def main():
    OUTPUT_TSV.write_text(
        "\n".join(["\t".join(HEADERS)] + ["\t".join(row) for row in ROWS]),
        encoding="utf-8",
    )

    doc = Document()
    section = doc.sections[0]
    section.orientation = WD_ORIENT.LANDSCAPE
    section.page_width, section.page_height = section.page_height, section.page_width
    for margin in (section.left_margin, section.right_margin, section.top_margin, section.bottom_margin):
        pass
    section.left_margin = Cm(1.5)
    section.right_margin = Cm(1.5)
    section.top_margin = Cm(1.5)
    section.bottom_margin = Cm(1.5)

    title = doc.add_paragraph()
    set_rtl(title, WD_ALIGN_PARAGRAPH.CENTER)
    r = title.add_run("جدول نمونه قواعد پرسشنامه تطبیقی")
    r.bold = True
    r.font.size = Pt(14)
    r.font.name = FONT

    table = doc.add_table(rows=1, cols=len(HEADERS))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    widths = [Cm(4.2), Cm(3.5), Cm(3.5), Cm(1.8), Cm(6.5)]
    for i, w in enumerate(widths):
        table.rows[0].cells[i].width = w

    for i, h in enumerate(HEADERS):
        write_cell(table.rows[0].cells[i], h, bold=True, center=True)
        shd = OxmlElement("w:shd")
        shd.set(qn("w:fill"), "D9E2F3")
        table.rows[0].cells[i]._tc.get_or_add_tcPr().append(shd)

    for row_data in ROWS:
        row = table.add_row()
        for i, val in enumerate(row_data):
            write_cell(row.cells[i], val, center=(i == 3))

    doc.save(OUTPUT_DOCX)
    print(f"Created: {OUTPUT_DOCX}")
    print(f"Created: {OUTPUT_TSV}")


if __name__ == "__main__":
    main()
