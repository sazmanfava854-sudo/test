#!/usr/bin/env python3
"""Compact matrix thesis table: 13 questions × 3 options in one readable table."""

import json
from pathlib import Path

from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

ROOT = Path(__file__).resolve().parents[1]
QUESTIONS_PATH = ROOT / "IoTRecommendation.Web" / "Data" / "Questions.json"
OUTPUT_DOCX = Path(__file__).resolve().parent / "thesis_questions_table_matrix.docx"
OUTPUT_TXT = Path(__file__).resolve().parent / "thesis_questions_table_matrix.txt"

FONT = "B Nazanin"
LABELS = ("الف", "ب", "ج")

CRITERION_FA = {
    "TransmissionRange": "برد انتقال",
    "LinkBudget": "بودجه پیوند",
    "EnergyConsumption": "مصرف انرژی",
    "HardwareCAPEX": "هزینه سرمایه‌ای",
    "AnnualConnectivityOPEX": "هزینه عملیاتی سالانه",
    "CellularSupport": "پشتیبانی سلولی",
    "DataRate": "نرخ داده",
    "RTTLatency": "تأخیر داده",
}

QUESTION_FA = {
    "Q01": "مساحت کل سایت استقرار چقدر است؟",
    "Q02": "شکل و توپوگرافی زمین چگونه است؟",
    "Q03": "چند مانع فیزیکی در محدوده وجود دارد؟",
    "Q04": "دسترسی به برق چگونه است؟",
    "Q05": "نصب دروازه یا اینترنت در محل چقدر آسان است؟",
    "Q06": "پوشش شبکه سلولی چگونه است؟",
    "Q07": "تعداد گره‌های حسگر چقدر است؟",
    "Q08": "میانگین فاصله بین گره‌ها چقدر است؟",
    "Q09": "حجم داده و فراوانی ارسال چگونه است؟",
    "Q10": "بودجه سرمایه‌ای سخت‌افزار چقدر است؟",
    "Q11": "تحمل هزینه عملیاتی سالانه چقدر است؟",
    "Q12": "عمر باتری مورد نیاز چقدر است؟",
    "Q13": "سرعت تحویل داده چقدر است؟",
}

OPTION_FA = {
    "Q01_A": "کوچک (<۱ هکتار)",
    "Q01_B": "متوسط (۱–۱۰ هکتار)",
    "Q01_C": "بزرگ (>۱۰ هکتار)",
    "Q02_A": "مسطح",
    "Q02_B": "کمی شیب‌دار",
    "Q02_C": "کوهستانی / ناهموار",
    "Q03_A": "کم (منطقه باز)",
    "Q03_B": "متوسط",
    "Q03_C": "زیاد (موانع متراکم)",
    "Q04_A": "برق شهری کامل",
    "Q04_B": "برق محدود",
    "Q04_C": "باتری / خورشیدی",
    "Q05_A": "آسان",
    "Q05_B": "با محدودیت",
    "Q05_C": "بسیار دشوار",
    "Q06_A": "پوشش خوب",
    "Q06_B": "پوشش متوسط",
    "Q06_C": "پوشش ضعیف",
    "Q07_A": "کم (۱–۱۰)",
    "Q07_B": "متوسط (۱۱–۱۰۰)",
    "Q07_C": "زیاد (>۱۰۰)",
    "Q08_A": "<۵۰ متر",
    "Q08_B": "۵۰–۵۰۰ متر",
    "Q08_C": ">۵۰۰ متر",
    "Q09_A": "کم",
    "Q09_B": "متوسط",
    "Q09_C": "زیاد",
    "Q10_A": "محدود",
    "Q10_B": "متوسط",
    "Q10_C": "منعطف",
    "Q11_A": "بسیار محدود",
    "Q11_B": "متوسط",
    "Q11_C": "منعطف",
    "Q12_A": "<۱ سال",
    "Q12_B": "۱–۳ سال",
    "Q12_C": ">۳ سال",
    "Q13_A": "غیربلادرنگ",
    "Q13_B": "نزدیک بلادرنگ",
    "Q13_C": "کاملاً بلادرنگ",
}

INTERPRETATIONS = {
    "Q01_A": "برد و پیوند اهمیت کمتر",
    "Q01_B": "بدون تغییر وزن",
    "Q01_C": "برد و پیوند اهمیت بیشتر",
    "Q02_A": "پیوند اهمیت کمتر",
    "Q02_B": "پیوند اهمیت بیشتر",
    "Q02_C": "پیوند اهمیت بیشتر",
    "Q03_A": "پیوند اهمیت کمتر",
    "Q03_B": "پیوند اهمیت بیشتر",
    "Q03_C": "پیوند اهمیت بیشتر",
    "Q04_A": "انرژی اهمیت کمتر",
    "Q04_B": "انرژی اهمیت بیشتر",
    "Q04_C": "انرژی اهمیت بیشتر",
    "Q05_A": "هزینه و سلولی اهمیت کمتر",
    "Q05_B": "سرمایه‌ای اهمیت بیشتر",
    "Q05_C": "هزینه و سلولی اهمیت بیشتر",
    "Q06_A": "سلولی اهمیت کمتر",
    "Q06_B": "سلولی اهمیت بیشتر",
    "Q06_C": "سلولی اهمیت کمتر",
    "Q07_A": "سرمایه‌ای اهمیت کمتر",
    "Q07_B": "بدون تغییر وزن",
    "Q07_C": "سرمایه‌ای و عملیاتی اهمیت بیشتر",
    "Q08_A": "برد اهمیت کمتر",
    "Q08_B": "بدون تغییر وزن",
    "Q08_C": "برد و پیوند اهمیت بیشتر",
    "Q09_A": "داده و تأخیر اهمیت کمتر",
    "Q09_B": "بدون تغییر وزن",
    "Q09_C": "داده، تأخیر و انرژی اهمیت بیشتر",
    "Q10_A": "سرمایه‌ای اهمیت بیشتر",
    "Q10_B": "بدون تغییر وزن",
    "Q10_C": "سرمایه‌ای اهمیت کمتر",
    "Q11_A": "عملیاتی اهمیت بیشتر",
    "Q11_B": "بدون تغییر وزن",
    "Q11_C": "عملیاتی اهمیت کمتر",
    "Q12_A": "انرژی اهمیت کمتر",
    "Q12_B": "بدون تغییر وزن",
    "Q12_C": "انرژی اهمیت بیشتر",
    "Q13_A": "تأخیر اهمیت کمتر",
    "Q13_B": "تأخیر اهمیت بیشتر",
    "Q13_C": "تأخیر اهمیت بیشتر",
}


def fmt_delta(delta: int) -> str:
    return f"+{delta}" if delta > 0 else f"−{abs(delta)}"


def format_option_block(option: dict) -> str:
    oid = option["id"]
    title = OPTION_FA.get(oid, option["text"])
    effects = option.get("effects") or []
    note = INTERPRETATIONS.get(oid, "")

    lines = [title]
    if not effects:
        lines.append("بدون اثر")
    else:
        parts = [
            f"{CRITERION_FA.get(e['criterionKey'], e['criterionKey'])} {fmt_delta(e['delta'])}"
            for e in effects
        ]
        lines.append("، ".join(parts))
    if note:
        lines.append(f"({note})")
    return "\n".join(lines)


def format_option_plain(label: str, option: dict) -> str:
    return f"{label}) {format_option_block(option).replace(chr(10), ' | ')}"


def set_rtl(paragraph, align=WD_ALIGN_PARAGRAPH.RIGHT):
    p_pr = paragraph._p.get_or_add_pPr()
    bidi = OxmlElement("w:bidi")
    bidi.set(qn("w:val"), "1")
    p_pr.append(bidi)
    paragraph.alignment = align


def set_cell_rtl(cell):
    tc_pr = cell._tc.get_or_add_tcPr()
    if not tc_pr.findall(qn("w:bidiVisual")):
        tc_pr.append(OxmlElement("w:bidiVisual"))
    v = OxmlElement("w:vAlign")
    v.set(qn("w:val"), "center")
    tc_pr.append(v)


def shade_cell(cell, fill: str):
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    shd.set(qn("w:val"), "clear")
    cell._tc.get_or_add_tcPr().append(shd)


def write_option_cell(cell, option: dict):
    set_cell_rtl(cell)
    p = cell.paragraphs[0]
    p.clear()
    set_rtl(p)

    oid = option["id"]
    title = OPTION_FA.get(oid, option["text"])
    effects = option.get("effects") or []
    note = INTERPRETATIONS.get(oid, "")

    r1 = p.add_run(f"{title}\n")
    r1.bold = True
    r1.font.name = FONT
    r1.font.size = Pt(9)

    if not effects:
        r2 = p.add_run("بدون اثر\n")
        r2.font.color.rgb = RGBColor(100, 100, 100)
    else:
        effect_text = "، ".join(
            f"{CRITERION_FA.get(e['criterionKey'], e['criterionKey'])} {fmt_delta(e['delta'])}"
            for e in effects
        )
        r2 = p.add_run(f"{effect_text}\n")
    r2.font.name = FONT
    r2.font.size = Pt(9)

    if note:
        r3 = p.add_run(note)
        r3.italic = True
        r3.font.name = FONT
        r3.font.size = Pt(8)
        r3.font.color.rgb = RGBColor(70, 70, 70)


def write_simple_cell(cell, text: str, *, bold=False, center=False, size=10):
    set_cell_rtl(cell)
    p = cell.paragraphs[0]
    p.clear()
    set_rtl(p, WD_ALIGN_PARAGRAPH.CENTER if center else WD_ALIGN_PARAGRAPH.RIGHT)
    run = p.add_run(text)
    run.bold = bold
    run.font.name = FONT
    run.font.size = Pt(size)


def build_plain_text(questions: list) -> str:
    lines = [
        "جدول فشرده پرسشنامه تطبیقی (۱۳ سؤال)",
        "=" * 60,
        "",
    ]
    for q in sorted(questions, key=lambda x: x["order"]):
        qid = q["id"]
        qtext = QUESTION_FA.get(qid, q["text"])
        lines.append(f"{q['order']}. {qtext}")
        for label, opt in zip(LABELS, q["options"]):
            lines.append(f"   {format_option_plain(label, opt)}")
        lines.append("")
    lines += [
        "راهنما: عدد مثبت (+۱) = افزایش اهمیت معیار | عدد منفی (−۱) = کاهش اهمیت معیار",
    ]
    return "\n".join(lines)


def main():
    questions = json.loads(QUESTIONS_PATH.read_text(encoding="utf-8"))
    OUTPUT_TXT.write_text(build_plain_text(questions), encoding="utf-8")

    doc = Document()
    section = doc.sections[0]
    section.orientation = WD_ORIENT.LANDSCAPE
    section.page_width, section.page_height = section.page_height, section.page_width
    for attr in ("left_margin", "right_margin", "top_margin", "bottom_margin"):
        setattr(section, attr, Cm(1.2))

    title = doc.add_paragraph()
    set_rtl(title, WD_ALIGN_PARAGRAPH.CENTER)
    tr = title.add_run("جدول فشرده پرسشنامه تطبیقی")
    tr.bold = True
    tr.font.size = Pt(14)
    tr.font.name = FONT

    sub = doc.add_paragraph()
    set_rtl(sub, WD_ALIGN_PARAGRAPH.CENTER)
    sr = sub.add_run("۱۳ سؤال — ۳۹ گزینه — نمایش ماتریسی")
    sr.font.size = Pt(11)
    sr.font.name = FONT

    doc.add_paragraph()

    headers = ["ردیف", "پرسش", "گزینه الف", "گزینه ب", "گزینه ج"]
    widths = [Cm(0.9), Cm(4.2), Cm(4.5), Cm(4.5), Cm(4.5)]

    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    for i, (h, w) in enumerate(zip(headers, widths)):
        write_simple_cell(table.rows[0].cells[i], h, bold=True, center=True, size=11)
        table.rows[0].cells[i].width = w
        shade_cell(table.rows[0].cells[i], "2F5496")
        for p in table.rows[0].cells[i].paragraphs:
            for r in p.runs:
                r.font.color.rgb = RGBColor(255, 255, 255)

    for idx, question in enumerate(sorted(questions, key=lambda q: q["order"]), 1):
        row = table.add_row()
        qid = question["id"]
        fill = "F2F2F2" if idx % 2 == 0 else "FFFFFF"

        write_simple_cell(row.cells[0], str(idx), center=True, size=10)
        write_simple_cell(row.cells[1], QUESTION_FA.get(qid, question["text"]), size=10)

        for col, option in enumerate(question["options"], 2):
            write_option_cell(row.cells[col], option)

        for cell in row.cells:
            shade_cell(cell, fill)

    note = doc.add_paragraph()
    set_rtl(note)
    nr = note.add_run(
        "راهنما: در هر گزینه، اعداد نشان‌دهنده تغییر اهمیت معیار در وزن‌دهی تطبیقی هستند "
        "(+۱ افزایش، −۱ کاهش). عبارت ایتالیک خلاصه تفسیر همان گزینه است."
    )
    nr.font.size = Pt(9)
    nr.font.name = FONT
    nr.italic = True

    doc.save(OUTPUT_DOCX)
    print(f"Created: {OUTPUT_DOCX}")
    print(f"Created: {OUTPUT_TXT}")


if __name__ == "__main__":
    main()
