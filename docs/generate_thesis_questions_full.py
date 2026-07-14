#!/usr/bin/env python3
"""Generate full Persian thesis table: 13 questions × 3 options."""

import json
from pathlib import Path

from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt

ROOT = Path(__file__).resolve().parents[1]
QUESTIONS_PATH = ROOT / "IoTRecommendation.Web" / "Data" / "Questions.json"
OUTPUT_DOCX = Path(__file__).resolve().parent / "thesis_questions_table_full.docx"
OUTPUT_TSV = Path(__file__).resolve().parent / "thesis_questions_table_full.tsv"

FONT = "B Nazanin"
HEADERS = ["پرسش", "پاسخ کاربر", "معیار تحت تأثیر", "اثر قاعده", "تفسیر"]

CRITERION_FA = {
    "TransmissionRange": "برد انتقال",
    "LinkBudget": "بودجه پیوند",
    "EnergyConsumption": "مصرف انرژی",
    "HardwareCAPEX": "هزینه سرمایه‌ای سخت‌افزار",
    "AnnualConnectivityOPEX": "هزینه عملیاتی سالانه اتصال",
    "CellularSupport": "پشتیبانی شبکه سلولی",
    "DataRate": "نرخ داده",
    "RTTLatency": "تأخیر انتقال داده",
}

QUESTION_FA = {
    "Q01": "مساحت کل سایت استقرار چقدر است؟",
    "Q02": "شکل و توپوگرافی زمین محل استقرار چگونه است؟",
    "Q03": "چند مانع فیزیکی در محدوده استقرار وجود دارد؟",
    "Q04": "دسترسی به برق در محل استقرار چگونه است؟",
    "Q05": "نصب دروازه یا نقطه دسترسی اینترنت در محل چقدر آسان است؟",
    "Q06": "پوشش شبکه سلولی در محل پروژه چگونه است؟",
    "Q07": "چند گره حسگر مستقر خواهد شد؟",
    "Q08": "میانگین فاصله بین گره‌های حسگر چقدر است؟",
    "Q09": "حجم داده و فراوانی گزارش‌دهی چگونه است؟",
    "Q10": "بودجه اولیه سرمایه‌ای سخت‌افزار به ازای هر گره چقدر است؟",
    "Q11": "تحمل هزینه عملیاتی سالانه چقدر است؟",
    "Q12": "انتظار عمر باتری مورد نیاز به ازای هر گره چقدر است؟",
    "Q13": "سرعت تحویل داده مورد نیاز چقدر است؟",
}

OPTION_FA = {
    "Q01_A": "کوچک (کمتر از ۱ هکتار)",
    "Q01_B": "متوسط (۱ تا ۱۰ هکتار)",
    "Q01_C": "بزرگ (بیش از ۱۰ هکتار)",
    "Q02_A": "مسطح",
    "Q02_B": "کمی شیب‌دار",
    "Q02_C": "کوهستانی و ناهموار",
    "Q03_A": "کم (منطقه باز)",
    "Q03_B": "متوسط",
    "Q03_C": "زیاد (موانع متراکم)",
    "Q04_A": "دسترسی کامل به شبکه برق",
    "Q04_B": "دسترسی محدود به برق",
    "Q04_C": "بدون دسترسی به برق (باتری یا خورشیدی)",
    "Q05_A": "آسان",
    "Q05_B": "با محدودیت",
    "Q05_C": "بسیار دشوار",
    "Q06_A": "پوشش خوب و استفاده مجاز",
    "Q06_B": "پوشش متوسط یا محدودیت جزئی",
    "Q06_C": "پوشش ضعیف یا عدم امکان استفاده از سلولی",
    "Q07_A": "کم (۱ تا ۱۰ گره)",
    "Q07_B": "متوسط (۱۱ تا ۱۰۰ گره)",
    "Q07_C": "زیاد (بیش از ۱۰۰ گره)",
    "Q08_A": "کمتر از ۵۰ متر",
    "Q08_B": "۵۰ تا ۵۰۰ متر",
    "Q08_C": "بیش از ۵۰۰ متر",
    "Q09_A": "کم (پیام‌های کوچک و کم‌تکرار)",
    "Q09_B": "متوسط",
    "Q09_C": "زیاد (پیام‌های بزرگ و پرتکرار)",
    "Q10_A": "بودجه محدود",
    "Q10_B": "بودجه متوسط",
    "Q10_C": "بودجه منعطف و مثمر",
    "Q11_A": "بسیار محدود",
    "Q11_B": "متوسط",
    "Q11_C": "منعطف",
    "Q12_A": "کمتر از ۱ سال",
    "Q12_B": "۱ تا ۳ سال",
    "Q12_C": "بیش از ۳ سال",
    "Q13_A": "غیربلادرنگ (تأخیر چند دقیقه‌ای قابل قبول)",
    "Q13_B": "نزدیک به بلادرنگ",
    "Q13_C": "کاملاً بلادرنگ",
}

INTERPRETATIONS = {
    "Q01_A": "سایت کوچک است؛ پوشش وسیع و بودجه پیوند بالا در اولویت طراحی شبکه نیست.",
    "Q01_B": "سایت با مقیاس متوسط است و این پاسخ تغییری در وزن معیارها ایجاد نمی‌کند.",
    "Q01_C": "سایت بزرگ است؛ برد انتقال و بودجه پیوند اهمیت بیشتری می‌یابند.",
    "Q02_A": "زمین مسطح تلفات سیگنال کمتری دارد؛ نیاز به بودجه پیوند بالا کاهش می‌یابد.",
    "Q02_B": "شیب جزئی ممکن است مسیر دید را محدود کند؛ بودجه پیوند اهمیت بیشتری پیدا می‌کند.",
    "Q02_C": "زمین ناهموار تلفات و چندمسیره‌سازی را افزایش می‌دهد؛ بودجه پیوند برجسته می‌شود.",
    "Q03_A": "منطقه باز موانع کمی دارد؛ نیاز به بودجه پیوند بالا کمتر است.",
    "Q03_B": "موانع متوسط احتمال تضعیف سیگنال را افزایش می‌دهد.",
    "Q03_C": "موانع متراکم نیاز به بودجه پیوند بالاتر را برجسته می‌کند.",
    "Q04_A": "دسترسی پایدار به برق، محدودیت انرژی را کاهش می‌دهد.",
    "Q04_B": "دسترسی محدود به برق، اهمیت مصرف انرژی پایین را افزایش می‌دهد.",
    "Q04_C": "وابستگی به باتری یا خورشیدی، مصرف انرژی را به معیار کلیدی تبدیل می‌کند.",
    "Q05_A": "نصب آسان دروازه، هزینه‌های سرمایه‌ای، عملیاتی و وابستگی به سلولی را کاهش می‌دهد.",
    "Q05_B": "محدودیت در نصب، هزینه سرمایه‌ای سخت‌افزار را افزایش می‌دهد.",
    "Q05_C": "نصب دشوار، هزینه‌ها و نیاز به راهکارهای جایگزین را برجسته می‌کند.",
    "Q06_A": "پوشش سلولی مناسب است؛ وابستگی تصمیم به معیار پشتیبانی سلولی کاهش می‌یابد.",
    "Q06_B": "پوشش متوسط، نیاز به ارزیابی دقیق‌تر فناوری‌های سلولی را افزایش می‌دهد.",
    "Q06_C": "پوشش ضعیف، فناوری‌های غیرسلولی یا زیرساخت اختصاصی اهمیت بیشتری می‌یابند.",
    "Q07_A": "تعداد کم گره، هزینه سرمایه‌ای کل را کاهش می‌دهد.",
    "Q07_B": "مقیاس متوسط است و تغییری در وزن معیارها ایجاد نمی‌کند.",
    "Q07_C": "استقرار گسترده، هزینه سرمایه‌ای و عملیاتی سالانه را برجسته می‌کند.",
    "Q08_A": "فاصله کوتاه بین گره‌ها، نیاز به برد انتقال بلند را کاهش می‌دهد.",
    "Q08_B": "فاصله متوسط است و تغییری در وزن معیارها ایجاد نمی‌کند.",
    "Q08_C": "فاصله زیاد، برد انتقال و بودجه پیوند را به معیارهای حیاتی تبدیل می‌کند.",
    "Q09_A": "داده کم و کم‌تکرار، اهمیت نرخ داده، تأخیر و مصرف انرژی را کاهش می‌دهد.",
    "Q09_B": "بار داده متوسط است و تغییری در وزن معیارها ایجاد نمی‌کند.",
    "Q09_C": "داده پرحجم و پرتکرار، نرخ داده، تأخیر و مصرف انرژی را برجسته می‌کند.",
    "Q10_A": "بودجه محدود، حساسیت به هزینه سرمایه‌ای سخت‌افزار را افزایش می‌دهد.",
    "Q10_B": "بودجه متوسط است و تغییری در وزن معیارها ایجاد نمی‌کند.",
    "Q10_C": "بودجه منعطف، محدودیت هزینه سرمایه‌ای را اهمیت کمتری می‌بخشد.",
    "Q11_A": "تحمل کم هزینه عملیاتی، اهمیت هزینه سالانه اتصال را افزایش می‌دهد.",
    "Q11_B": "تحمل متوسط است و تغییری در وزن معیارها ایجاد نمی‌کند.",
    "Q11_C": "تحمل منعطف، محدودیت هزینه عملیاتی را اهمیت کمتری می‌بخشد.",
    "Q12_A": "عمر باتری کوتاه، حساسیت به مصرف انرژی را کاهش می‌دهد.",
    "Q12_B": "عمر باتری متوسط است و تغییری در وزن معیارها ایجاد نمی‌کند.",
    "Q12_C": "عمر باتری بلند، فناوری‌های کم‌مصرف اهمیت بیشتری می‌یابند.",
    "Q13_A": "تحویل غیربلادرنگ، اهمیت تأخیر انتقال داده را کاهش می‌دهد.",
    "Q13_B": "نیاز نزدیک به بلادرنگ، اهمیت تأخیر را افزایش می‌دهد.",
    "Q13_C": "نیاز کاملاً بلادرنگ، تأخیر انتقال داده را به معیار بحرانی تبدیل می‌کند.",
}


def fmt_delta(delta: int) -> str:
    return f"+{delta}" if delta > 0 else f"−{abs(delta)}"


def build_rows(questions: list) -> list[tuple[str, str, str, str, str]]:
    """One row per answer option (13 questions × 3 options = 39 rows)."""
    rows = []
    for question in sorted(questions, key=lambda q: q["order"]):
        qid = question["id"]
        qtext = QUESTION_FA.get(qid, question["text"])
        for option in question["options"]:
            oid = option["id"]
            atext = OPTION_FA.get(oid, option["text"])
            interpretation = INTERPRETATIONS.get(oid, "")
            effects = option.get("effects") or []

            if not effects:
                rows.append((qtext, atext, "—", "—", interpretation))
                continue

            criteria = [CRITERION_FA.get(e["criterionKey"], e["criterionKey"]) for e in effects]
            deltas = [fmt_delta(e["delta"]) for e in effects]
            rows.append((qtext, atext, "؛ ".join(criteria), "؛ ".join(deltas), interpretation))

    return rows


def set_rtl(paragraph, align=WD_ALIGN_PARAGRAPH.RIGHT):
    p_pr = paragraph._p.get_or_add_pPr()
    bidi = OxmlElement("w:bidi")
    bidi.set(qn("w:val"), "1")
    p_pr.append(bidi)
    paragraph.alignment = align


def set_cell_rtl(cell):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_pr.append(OxmlElement("w:bidiVisual"))
    v = OxmlElement("w:vAlign")
    v.set(qn("w:val"), "center")
    tc_pr.append(v)
    for p in cell.paragraphs:
        set_rtl(p)


def write_cell(cell, text, *, bold=False, center=False):
    set_cell_rtl(cell)
    p = cell.paragraphs[0]
    p.clear()
    set_rtl(p, WD_ALIGN_PARAGRAPH.CENTER if center else WD_ALIGN_PARAGRAPH.RIGHT)
    run = p.add_run(text)
    run.bold = bold
    run.font.name = FONT
    run.font.size = Pt(10)


def main():
    questions = json.loads(QUESTIONS_PATH.read_text(encoding="utf-8"))
    rows = build_rows(questions)

    OUTPUT_TSV.write_text(
        "\n".join(["\t".join(HEADERS)] + ["\t".join(r) for r in rows]),
        encoding="utf-8",
    )

    doc = Document()
    section = doc.sections[0]
    section.orientation = WD_ORIENT.LANDSCAPE
    section.page_width, section.page_height = section.page_height, section.page_width
    section.left_margin = Cm(1.2)
    section.right_margin = Cm(1.2)
    section.top_margin = Cm(1.2)
    section.bottom_margin = Cm(1.2)

    title = doc.add_paragraph()
    set_rtl(title, WD_ALIGN_PARAGRAPH.CENTER)
    r = title.add_run("جدول کامل قواعد پرسشنامه تطبیقی (۱۳ پرسش)")
    r.bold = True
    r.font.size = Pt(14)
    r.font.name = FONT

    table = doc.add_table(rows=1, cols=len(HEADERS))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    widths = [Cm(4.0), Cm(3.2), Cm(3.0), Cm(1.5), Cm(6.0)]
    for i, w in enumerate(widths):
        table.rows[0].cells[i].width = w

    for i, h in enumerate(HEADERS):
        write_cell(table.rows[0].cells[i], h, bold=True, center=True)
        shd = OxmlElement("w:shd")
        shd.set(qn("w:fill"), "D9E2F3")
        table.rows[0].cells[i]._tc.get_or_add_tcPr().append(shd)

    for row_data in rows:
        row = table.add_row()
        for i, val in enumerate(row_data):
            write_cell(row.cells[i], val, center=(i == 3))

    note = doc.add_paragraph()
    set_rtl(note)
    nr = note.add_run(
        "توضیح: این جدول شامل ۱۳ پرسش و ۳۹ گزینه پاسخ است. "
        "اگر یک گزینه بر چند معیار اثر بگذارد، نام معیارها و اثر قاعده با «؛» از هم جدا شده‌اند. "
        "علامت «—» به معنای عدم تأثیر بر معیارها است."
    )
    nr.font.size = Pt(9)
    nr.font.name = FONT
    nr.italic = True

    doc.save(OUTPUT_DOCX)
    print(f"Rows: {len(rows)}")
    print(f"Created: {OUTPUT_DOCX}")
    print(f"Created: {OUTPUT_TSV}")


if __name__ == "__main__":
    main()
