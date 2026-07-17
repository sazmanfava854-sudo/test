#!/usr/bin/env python3
"""Generate Word document for thesis section 4-5-2 VIKOR."""

from io import BytesIO
from pathlib import Path

import matplotlib.pyplot as plt
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Cm, Pt

OUTPUT = Path("/opt/cursor/artifacts/4-5-2_VIKOR_Ranking.docx")
OUTPUT.parent.mkdir(parents=True, exist_ok=True)


def render_formula(latex: str, fontsize: int = 17) -> BytesIO:
    fig = plt.figure(figsize=(9, 1.0))
    fig.patch.set_alpha(0.0)
    fig.text(0.5, 0.5, f"${latex}$", fontsize=fontsize, ha="center", va="center")
    buf = BytesIO()
    fig.savefig(buf, format="png", dpi=220, bbox_inches="tight", transparent=True)
    plt.close(fig)
    buf.seek(0)
    return buf


def set_rtl(paragraph):
    p_pr = paragraph._p.get_or_add_pPr()
    bidi = p_pr.find(qn("w:bidi"))
    if bidi is None:
        bidi = p_pr.makeelement(qn("w:bidi"), {})
        p_pr.append(bidi)
    bidi.set(qn("w:val"), "1")


def add_para(doc, text, bold=False, size=14, center=False):
    p = doc.add_paragraph()
    set_rtl(p)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER if center else WD_ALIGN_PARAGRAPH.JUSTIFY
    run = p.add_run(text)
    run.bold = bold
    run.font.size = Pt(size)
    run.font.name = "B Nazanin"
    r = run._element.rPr
    if r is not None:
        r.rFonts.set(qn("w:ascii"), "Times New Roman")
        r.rFonts.set(qn("w:hAnsi"), "Times New Roman")
        r.rFonts.set(qn("w:cs"), "B Nazanin")
    return p


def add_formula(doc, latex: str, caption: str = ""):
    if caption:
        add_para(doc, caption, bold=True, size=13, center=True)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run().add_picture(render_formula(latex), width=Cm(14.5))


def build():
    doc = Document()

    add_para(doc, "۴-۵-۲- رتبه‌بندی با VIKOR", bold=True, size=16, center=True)

    body = [
        (
            "روش VIKOR یکی از روش‌های تصمیم‌گیری چندمعیاره است که برای رتبه‌بندی گزینه‌ها "
            "در شرایط وجود معیارهای متعارض به‌کار می‌رود. منطق اصلی این روش بر شناسایی یک "
            "راه‌حل توافقی استوار است؛ به این معنا که گزینه منتخب باید از یک‌سو بیشترین "
            "مطلوبیت کلی را برای مجموعه معیارها فراهم کند و از سوی دیگر، بیشترین نارضایتی "
            "یا ضعف بحرانی را به حداقل برساند. به همین دلیل، این روش تنها به عملکرد کلی "
            "گزینه‌ها توجه ندارد، بلکه بدترین عملکرد هر گزینه را نیز در فرآیند ارزیابی وارد "
            "می‌کند."
        ),
        (
            "در روش VIKOR، ابتدا بهترین و بدترین مقدار هر معیار در میان گزینه‌های خوشه "
            "منتخب تعیین می‌شود. سپس برای هر گزینه، دو شاخص اصلی محاسبه می‌گردد. شاخص "
            "نخست، شاخص سودمندی گروهی با نماد S_i است که فاصله کلی گزینه از وضعیت ایده‌آل "
            "را نشان می‌دهد و از رابطه (۴-۱) به‌دست می‌آید:"
        ),
    ]
    for t in body:
        add_para(doc, t)

    add_formula(doc, r"S_i=\sum_{j=1}^{n} w_j \frac{f_j^* - f_{ij}}{f_j^* - f_j^-}", "رابطه (۴-۱)")

    add_para(
        doc,
        "در این رابطه، w_j وزن معیار jام، f_ij عملکرد گزینه iام در معیار jام، f_j^* بهترین "
        "مقدار معیار و f_j^- بدترین مقدار آن است. برای معیارهای سود، f_j^* برابر بیشینه و "
        "f_j^- برابر کمینه مقادیر گزینه‌هاست؛ و برای معیارهای هزینه، این تعریف معکوس "
        "می‌شود. هرچه مقدار S_i کمتر باشد، گزینه از نظر عملکرد کلی به وضعیت مطلوب نزدیک‌تر "
        "است.",
    )

    add_para(
        doc,
        "شاخص دوم، شاخص بیشترین پشیمانی با نماد R_i است که بر بدترین عملکرد گزینه تمرکز "
        "دارد و از رابطه (۴-۲) محاسبه می‌شود:",
    )

    add_formula(
        doc,
        r"R_i=\max_{j}\left\{w_j \frac{f_j^* - f_{ij}}{f_j^* - f_j^-}\right\}",
        "رابطه (۴-۲)",
    )

    add_para(
        doc,
        "این شاخص نشان می‌دهد گزینه موردنظر در نامطلوب‌ترین معیار خود چه میزان از وضعیت "
        "ایده‌آل فاصله دارد. در نتیجه، هرچه مقدار R_i کوچک‌تر باشد، گزینه از نظر کنترل "
        "نقاط ضعف بحرانی وضعیت مناسب‌تری خواهد داشت.",
    )

    add_para(
        doc,
        "در مرحله بعد، شاخص نهایی VIKOR با نماد Q_i برای هر گزینه محاسبه می‌شود تا "
        "رتبه‌بندی نهایی بر اساس آن انجام گیرد (رابطه ۴-۳):",
    )

    add_formula(
        doc,
        r"Q_i = v\frac{S_i-S^*}{S^- - S^*} + (1-v)\frac{R_i-R^*}{R^- - R^*}",
        "رابطه (۴-۳)",
    )

    add_para(
        doc,
        "در این رابطه، S^* و S^- به‌ترتیب کمینه و بیشینه S_i، و R^* و R^- به‌ترتیب کمینه و "
        "بیشینه R_i در میان گزینه‌ها هستند. پارامتر v (بازه ۰ تا ۱) وزن راهبرد سودمندی "
        "گروهی در برابر راهبرد حداقل‌سازی پشیمانی است. در این پژوهش، مقدار v از فایل "
        "تنظیمات سامانه (Settings.json) خوانده شده و برای ارزیابی اصلی برابر ۰٫۵ در نظر "
        "گرفته شده است؛ بنابراین در رتبه‌بندی نهایی، اهمیت یکسانی به عملکرد کلی گزینه‌ها "
        "و پرهیز از ضعف شدید در یک معیار خاص داده شده است. در VIKOR، گزینه‌ای مطلوب‌تر "
        "است که مقدار Q_i کمتری داشته باشد.",
    )

    add_para(
        doc,
        "پیاده‌سازی این مرحله در سامانه IoTRecommendation و در کلاس VikorCalculator انجام "
        "شده است. پس از محاسبه S_i، R_i و Q_i، گزینه‌ها به‌صورت صعودی بر اساس Q_i مرتب "
        "می‌شوند و دو شرط پذیرش راه‌حل توافقی (C1 و C2) بررسی می‌گردد. نتایج عددی "
        "رتبه‌بندی در جدول (۴-XX) گزارش شده است.",
    )

    add_para(
        doc,
        "بر این اساس، روش VIKOR امکان رتبه‌بندی گزینه‌ها را با تکیه بر یک منطق متوازن "
        "فراهم می‌کند و گزینه‌ای را در اولویت قرار می‌دهد که در عین برخورداری از مطلوبیت "
        "کلی مناسب، از نظر نقاط ضعف بحرانی نیز در سطح قابل‌قبولی قرار داشته باشد.",
    )

    doc.add_page_break()
    add_para(doc, "پیوست: نحوه درج صحیح فرمول در Microsoft Word", bold=True, size=15, center=True)

    guide = [
        "۱. کلیدهای Alt + = را بزنید تا کادر معادله باز شود.",
        "۲. در Word ۲۰۱۹/Microsoft 365: از تب Equation (معادله) گزینه Convert و سپس LaTeX را انتخاب کنید.",
        "۳. کد LaTeX را Paste کنید و Enter بزنید.",
        "۴. سپس Convert → Professional را بزنید تا فرمول نهایی ساخته شود.",
        "",
        "——— اگر LaTeX فعال نیست، از فرمت UnicodeMath استفاده کنید ———",
        "پس از Paste در کادر معادله، بین بخش‌ها Space بزنید تا Word فرمول را بسازد.",
        "",
        "S_i=\\sum_(j=1)^n w_j ((f_j^*-f_ij))/(f_j^*-f_j^-)",
        "",
        "R_i=\\max_j(w_j ((f_j^*-f_ij))/(f_j^*-f_j^-))",
        "",
        "Q_i=v ((S_i-S^*))/(S^- -S^*)+(1-v) ((R_i-R^*))/(R^- -R^*)",
        "",
        "نکات:",
        "• فرمول را فقط داخل کادر معادله Paste کنید، نه در متن عادی.",
        "• اگر نمایش به‌هم‌ریخته شد، معادله را انتخاب کنید و Linear را به Professional تبدیل کنید.",
        "• برای زیرنویس از _ و برای توان از ^ استفاده کنید.",
    ]
    for line in guide:
        if line.startswith("S_i") or line.startswith("R_i") or line.startswith("Q_i"):
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            r = p.add_run(line)
            r.font.name = "Cambria Math"
            r.font.size = Pt(12)
        else:
            add_para(doc, line, size=12)

    doc.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    build()
