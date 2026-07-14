#!/usr/bin/env python3
"""Generate sample thesis table with a few representative questions only."""

import json
from pathlib import Path

from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

ROOT = Path(__file__).resolve().parents[1]
QUESTIONS_PATH = ROOT / "IoTRecommendation.Web" / "Data" / "Questions.json"
OUTPUT_DOCX = Path(__file__).resolve().parent / "thesis_questions_table_sample.docx"
OUTPUT_TSV = Path(__file__).resolve().parent / "thesis_questions_table_sample.tsv"

# Representative questions for thesis main text
SAMPLE_IDS = ["Q01", "Q04", "Q09", "Q10"]

FONT = "B Nazanin"

# Import shared data from matrix generator
import importlib.util

spec = importlib.util.spec_from_file_location("matrix", Path(__file__).parent / "generate_thesis_questions_matrix.py")
matrix = importlib.util.module_from_spec(spec)
spec.loader.exec_module(matrix)


def cell_text(option: dict) -> str:
    oid = option["id"]
    title = matrix.OPTION_FA.get(oid, option["text"])
    effects = option.get("effects") or []
    note = matrix.INTERPRETATIONS.get(oid, "")
    if not effects:
        mid = "بدون اثر"
    else:
        mid = "، ".join(
            f"{matrix.CRITERION_FA.get(e['criterionKey'], e['criterionKey'])} {matrix.fmt_delta(e['delta'])}"
            for e in effects
        )
    return f"{title} | {mid} | ({note})" if note else f"{title} | {mid}"


def set_rtl(paragraph, align=WD_ALIGN_PARAGRAPH.RIGHT):
    p_pr = paragraph._p.get_or_add_pPr()
    bidi = OxmlElement("w:bidi")
    bidi.set(qn("w:val"), "1")
    p_pr.append(bidi)
    paragraph.alignment = align


def set_cell_rtl(cell):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_pr.append(OxmlElement("w:bidiVisual"))
    v = OxmlElement("w:vAlign")
    v.set(qn("w:val"), "center")
    tc_pr.append(v)


def shade_cell(cell, fill: str):
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    shd.set(qn("w:val"), "clear")
    cell._tc.get_or_add_tcPr().append(shd)


def write_option_cell(cell, option: dict):
    set_cell_rtl(cell)
    p = cell.paragraphs[0]
    p.clear()
    set_rtl(p)
    oid = option["id"]
    title = matrix.OPTION_FA.get(oid, option["text"])
    effects = option.get("effects") or []
    note = matrix.INTERPRETATIONS.get(oid, "")

    r1 = p.add_run(f"{title}\n")
    r1.bold = True
    r1.font.name = FONT
    r1.font.size = Pt(9)

    if not effects:
        r2 = p.add_run("بدون اثر\n")
        r2.font.color.rgb = RGBColor(100, 100, 100)
    else:
        effect_text = "، ".join(
            f"{matrix.CRITERION_FA.get(e['criterionKey'], e['criterionKey'])} {matrix.fmt_delta(e['delta'])}"
            for e in effects
        )
        r2 = p.add_run(f"{effect_text}\n")
    r2.font.name = FONT
    r2.font.size = Pt(9)

    if note:
        r3 = p.add_run(note)
        r3.italic = True
        r3.font.name = FONT
        r3.font.size = Pt(8)
        r3.font.color.rgb = RGBColor(70, 70, 70)


def write_simple_cell(cell, text: str, *, bold=False, center=False, size=10, fill=None):
    set_cell_rtl(cell)
    p = cell.paragraphs[0]
    p.clear()
    set_rtl(p, WD_ALIGN_PARAGRAPH.CENTER if center else WD_ALIGN_PARAGRAPH.RIGHT)
    run = p.add_run(text)
    run.bold = bold
    run.font.name = FONT
    run.font.size = Pt(size)
    if fill:
        shade_cell(cell, fill)


def main():
    all_questions = {q["id"]: q for q in json.loads(QUESTIONS_PATH.read_text(encoding="utf-8"))}
    questions = [all_questions[qid] for qid in SAMPLE_IDS]

    headers = ["ردیف", "پرسش", "گزینه الف", "گزینه ب", "گزینه ج"]
    rows = [headers]
    for i, q in enumerate(questions, 1):
        opts = q["options"]
        rows.append([
            str(i),
            matrix.QUESTION_FA.get(q["id"], q["text"]),
            cell_text(opts[0]),
            cell_text(opts[1]),
            cell_text(opts[2]),
        ])
    OUTPUT_TSV.write_text("\n".join("\t".join(r) for r in rows), encoding="utf-8")

    doc = Document()
    section = doc.sections[0]
    section.orientation = WD_ORIENT.LANDSCAPE
    section.page_width, section.page_height = section.page_height, section.page_width
    for attr in ("left_margin", "right_margin", "top_margin", "bottom_margin"):
        setattr(section, attr, Cm(1.5))

    title = doc.add_paragraph()
    set_rtl(title, WD_ALIGN_PARAGRAPH.CENTER)
    tr = title.add_run("جدول نمونه قواعد پرسشنامه تطبیقی")
    tr.bold = True
    tr.font.size = Pt(14)
    tr.font.name = FONT

    sub = doc.add_paragraph()
    set_rtl(sub, WD_ALIGN_PARAGRAPH.CENTER)
    sr = sub.add_run("(نمونه‌ای از ۴ پرسش — جدول کامل در پیوست)")
    sr.font.size = Pt(11)
    sr.font.name = FONT
    sr.italic = True

    doc.add_paragraph()

    table = doc.add_table(rows=1, cols=5)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    widths = [Cm(0.9), Cm(4.0), Cm(4.5), Cm(4.5), Cm(4.5)]

    for i, h in enumerate(headers):
        write_simple_cell(table.rows[0].cells[i], h, bold=True, center=True, size=11, fill="2F5496")
        table.rows[0].cells[i].width = widths[i]
        for p in table.rows[0].cells[i].paragraphs:
            for r in p.runs:
                r.font.color.rgb = RGBColor(255, 255, 255)

    for i, q in enumerate(questions, 1):
        row = table.add_row()
        fill = "F2F2F2" if i % 2 == 0 else "FFFFFF"
        write_simple_cell(row.cells[0], str(i), center=True, size=10)
        write_simple_cell(row.cells[1], matrix.QUESTION_FA.get(q["id"], q["text"]), size=10)
        for col, opt in enumerate(q["options"], 2):
            write_option_cell(row.cells[col], opt)
            shade_cell(row.cells[col], fill)
        shade_cell(row.cells[0], fill)
        shade_cell(row.cells[1], fill)

    note = doc.add_paragraph()
    set_rtl(note)
    nr = note.add_run(
        "توضیح: این جدول تنها نمونه‌ای از قواعد تطبیقی است. "
        "مقادیر +۱ و −۱ نشان‌دهنده افزایش یا کاهش اهمیت نسبی معیار در وزن‌دهی تطبیقی می‌باشند. "
        "جدول کامل ۱۳ پرسش در پیوست آورده شده است."
    )
    nr.font.size = Pt(9)
    nr.font.name = FONT
    nr.italic = True

    doc.save(OUTPUT_DOCX)
    print(f"Created: {OUTPUT_DOCX}")
    print(f"Created: {OUTPUT_TSV}")


if __name__ == "__main__":
    main()
