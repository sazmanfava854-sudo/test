#!/usr/bin/env python3
"""Generate thesis Word table for questionnaire questions section (Persian only)."""

import json
from pathlib import Path

from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt

ROOT = Path(__file__).resolve().parents[1]
QUESTIONS_PATH = ROOT / "IoTRecommendation.Web" / "Data" / "Questions.json"
OUTPUT_PATH = ROOT / "docs" / "thesis_questions_table.docx"

FONT_NAME = "B Nazanin"
FONT_SIZE_BODY = Pt(11)
FONT_SIZE_HEADER = Pt(12)
FONT_SIZE_TITLE = Pt(14)

CRITERION_FA = {
    "TransmissionRange": "دامنه انتقال",
    "LinkBudget": "بودجه پیوند",
    "EnergyConsumption": "مصرف انرژی",
    "HardwareCAPEX": "هزینه سرمایه‌ای سخت‌افزار",
    "AnnualConnectivityOPEX": "هزینه عملیاتی سالانه اتصال",
    "CellularSupport": "پشتیبانی شبکه سلولی",
    "DataRate": "نرخ داده",
    "RTTLatency": "تأخیر انتقال داده",
}

QUESTION_FA = {
    "Q01": "مساحت کل زمین محل استقرار چقدر است؟",
    "Q02": "شکل و توپوگرافی زمین محل استقرار چگونه است؟",
    "Q03": "چند مانع فیزیکی در محدوده استقرار وجود دارد؟",
    "Q04": "دسترسی به برق در محل استقرار چگونه است؟",
    "Q05": "نصب دروازه یا نقطه دسترسی اینترنت در محل چقدر آسان است؟",
    "Q06": "پوشش شبکه سلولی در محل پروژه چگونه است؟",
    "Q07": "چند گره حسگر مستقر خواهد شد؟",
    "Q08": "میانگین فاصله بین گره‌های حسگر چقدر است؟",
    "Q09": "حجم داده و فراوانی گزارش‌دهی چگونه است؟",
    "Q10": "بودجه اولیه سرمایه‌ای سخت‌افزار به ازای هر گره چقدر است؟",
    "Q11": "تحمل هزینه عملیاتی سالانه چقدر است؟",
    "Q12": "انتظار عمر باتری مورد نیاز به ازای هر گره چقدر است؟",
    "Q13": "سرعت تحویل داده مورد نیاز چقدر است؟",
}

OPTION_FA = {
    "Q01_A": "کوچک (کمتر از ۱ هکتار)",
    "Q01_B": "متوسط (۱ تا ۱۰ هکتار)",
    "Q01_C": "بزرگ (بیش از ۱۰ هکتار)",
    "Q02_A": "مسطح",
    "Q02_B": "کمی شیب‌دار",
    "Q02_C": "کوهستانی و ناهموار",
    "Q03_A": "کم (منطقه باز)",
    "Q03_B": "متوسط",
    "Q03_C": "زیاد (موانع متراکم)",
    "Q04_A": "دسترسی کامل به شبکه برق",
    "Q04_B": "دسترسی محدود به برق",
    "Q04_C": "بدون دسترسی به برق (باتری یا خورشیدی)",
    "Q05_A": "آسان",
    "Q05_B": "با محدودیت",
    "Q05_C": "بسیار دشوار",
    "Q06_A": "پوشش خوب و استفاده مجاز",
    "Q06_B": "پوشش متوسط یا محدودیت جزئی",
    "Q06_C": "پوشش ضعیف یا عدم امکان استفاده از سلولی",
    "Q07_A": "کم (۱ تا ۱۰ گره)",
    "Q07_B": "متوسط (۱۱ تا ۱۰۰ گره)",
    "Q07_C": "زیاد (بیش از ۱۰۰ گره)",
    "Q08_A": "کمتر از ۵۰ متر",
    "Q08_B": "۵۰ تا ۵۰۰ متر",
    "Q08_C": "بیش از ۵۰۰ متر",
    "Q09_A": "کم (پیام‌های کوچک و کم‌تکرار)",
    "Q09_B": "متوسط",
    "Q09_C": "زیاد (پیام‌های بزرگ و پرتکرار)",
    "Q10_A": "بودجه محدود",
    "Q10_B": "بودجه متوسط",
    "Q10_C": "بودجه منعطف و مثمر",
    "Q11_A": "بسیار محدود",
    "Q11_B": "متوسط",
    "Q11_C": "منعطف",
    "Q12_A": "کمتر از ۱ سال",
    "Q12_B": "۱ تا ۳ سال",
    "Q12_C": "بیش از ۳ سال",
    "Q13_A": "غیربلادرنگ (تأخیر چند دقیقه‌ای قابل قبول)",
    "Q13_B": "نزدیک به بلادرنگ",
    "Q13_C": "کاملاً بلادرنگ",
}

INTERPRETATIONS = {
    ("Q01", "Q01_A"): "محل استقرار کوچک است؛ نیاز به دامنه انتقال بلند و بودجه پیوند بالا کمتر است و اهمیت این معیارها در وزن‌دهی تطبیقی کاهش می‌یابد.",
    ("Q01", "Q01_B"): "شرایط متوسط است و این پاسخ تغییری در امتیاز تعدیل معیارها ایجاد نمی‌کند.",
    ("Q01", "Q01_C"): "محل استقرار بزرگ است؛ نیاز به پوشش وسیع‌تر و پیوند قوی‌تر افزایش می‌یابد.",
    ("Q02", "Q02_A"): "زمین مسطح تلفات سیگنال کمتری دارد؛ اهمیت بودجه پیوند کاهش می‌یابد.",
    ("Q02", "Q02_B"): "شیب جزئی ممکن است مسیر دید را محدود کند؛ اهمیت بودجه پیوند افزایش می‌یابد.",
    ("Q02", "Q02_C"): "زمین ناهموار تلفات و چندمسیره‌سازی را افزایش می‌دهد؛ بودجه پیوند اهمیت بیشتری می‌یابد.",
    ("Q03", "Q03_A"): "منطقه باز موانع کمی دارد؛ نیاز به بودجه پیوند بالا کمتر است.",
    ("Q03", "Q03_B"): "موانع متوسط احتمال تضعیف سیگنال را افزایش می‌دهد.",
    ("Q03", "Q03_C"): "موانع متراکم نیاز به بودجه پیوند بالاتر و فناوری مقاوم‌تر را برجسته می‌کند.",
    ("Q04", "Q04_A"): "دسترسی کامل به برق، حساسیت به مصرف انرژی را کاهش می‌دهد.",
    ("Q04", "Q04_B"): "دسترسی محدود به برق، اهمیت مصرف انرژی پایین را افزایش می‌دهد.",
    ("Q04", "Q04_C"): "وابستگی به باتری یا خورشیدی، مصرف انرژی را به معیار کلیدی تبدیل می‌کند.",
    ("Q05", "Q05_A"): "نصب آسان دروازه، هزینه‌های سرمایه‌ای، عملیاتی و نیاز به شبکه سلولی را کاهش می‌دهد.",
    ("Q05", "Q05_B"): "محدودیت در نصب، هزینه سرمایه‌ای سخت‌افزار را افزایش می‌دهد.",
    ("Q05", "Q05_C"): "نصب دشوار، هزینه‌ها و وابستگی به راهکارهای جایگزین را برجسته می‌کند.",
    ("Q06", "Q06_A"): "پوشش سلولی مناسب است؛ اهمیت معیار پشتیبانی سلولی کاهش می‌یابد.",
    ("Q06", "Q06_B"): "پوشش متوسط، نیاز به ارزیابی دقیق‌تر فناوری‌های سلولی را افزایش می‌دهد.",
    ("Q06", "Q06_C"): "پوشش ضعیف است؛ فناوری‌های غیرسلولی یا زیرساخت اختصاصی اهمیت بیشتری می‌یابند.",
    ("Q07", "Q07_A"): "تعداد کم گره، هزینه سرمایه‌ای کل را کاهش می‌دهد.",
    ("Q07", "Q07_B"): "مقیاس متوسط است و تغییری در امتیاز تعدیل ایجاد نمی‌کند.",
    ("Q07", "Q07_C"): "استقرار گسترده، هزینه سرمایه‌ای و عملیاتی سالانه را برجسته می‌کند.",
    ("Q08", "Q08_A"): "فاصله کوتاه بین گره‌ها، نیاز به دامنه انتقال بلند را کاهش می‌دهد.",
    ("Q08", "Q08_B"): "فاصله متوسط است و تغییری در امتیاز تعدیل ایجاد نمی‌کند.",
    ("Q08", "Q08_C"): "فاصله زیاد، دامنه انتقال و بودجه پیوند را به معیارهای حیاتی تبدیل می‌کند.",
    ("Q09", "Q09_A"): "داده کم و کم‌تکرار، اهمیت نرخ داده، تأخیر و مصرف انرژی را کاهش می‌دهد.",
    ("Q09", "Q09_B"): "بار داده متوسط است و تغییری در امتیاز تعدیل ایجاد نمی‌کند.",
    ("Q09", "Q09_C"): "داده پرحجم و پرتکرار، نرخ داده، تأخیر و مصرف انرژی را برجسته می‌کند.",
    ("Q10", "Q10_A"): "بودجه محدود، حساسیت به هزینه سرمایه‌ای سخت‌افزار را افزایش می‌دهد.",
    ("Q10", "Q10_B"): "بودجه متوسط است و تغییری در امتیاز تعدیل ایجاد نمی‌کند.",
    ("Q10", "Q10_C"): "بودجه منعطف، محدودیت هزینه سرمایه‌ای را اهمیت کمتری می‌بخشد.",
    ("Q11", "Q11_A"): "تحمل کم هزینه عملیاتی، اهمیت هزینه سالانه اتصال را افزایش می‌دهد.",
    ("Q11", "Q11_B"): "تحمل متوسط است و تغییری در امتیاز تعدیل ایجاد نمی‌کند.",
    ("Q11", "Q11_C"): "تحمل منعطف، محدودیت هزینه عملیاتی را اهمیت کمتری می‌بخشد.",
    ("Q12", "Q12_A"): "عمر باتری کوتاه، حساسیت به مصرف انرژی را کاهش می‌دهد.",
    ("Q12", "Q12_B"): "عمر باتری متوسط است و تغییری در امتیاز تعدیل ایجاد نمی‌کند.",
    ("Q12", "Q12_C"): "عمر باتری بلند، فناوری‌های کم‌مصرف اهمیت بیشتری می‌یابند.",
    ("Q13", "Q13_A"): "تحویل غیربلادرنگ، اهمیت تأخیر انتقال داده را کاهش می‌دهد.",
    ("Q13", "Q13_B"): "نیاز نزدیک به بلادرنگ، اهمیت تأخیر را افزایش می‌دهد.",
    ("Q13", "Q13_C"): "نیاز کاملاً بلادرنگ، تأخیر انتقال داده را به معیار بحرانی تبدیل می‌کند.",
}


def set_paragraph_rtl(paragraph, align=WD_ALIGN_PARAGRAPH.RIGHT):
    p_pr = paragraph._p.get_or_add_pPr()
    bidi = OxmlElement("w:bidi")
    bidi.set(qn("w:val"), "1")
    p_pr.append(bidi)
    paragraph.alignment = align


def set_cell_rtl(cell, valign="center"):
    tc_pr = cell._tc.get_or_add_tcPr()
    if not tc_pr.findall(qn("w:bidiVisual")):
        tc_pr.append(OxmlElement("w:bidiVisual"))
    v_align = OxmlElement("w:vAlign")
    v_align.set(qn("w:val"), valign)
    tc_pr.append(v_align)
    for paragraph in cell.paragraphs:
        set_paragraph_rtl(paragraph)


def shade_cell(cell, fill="D9E2F3"):
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), fill)
    shading.set(qn("w:val"), "clear")
    cell._tc.get_or_add_tcPr().append(shading)


def write_cell_text(cell, text, *, bold=False, size=FONT_SIZE_BODY, align=WD_ALIGN_PARAGRAPH.RIGHT):
    set_cell_rtl(cell)
    p = cell.paragraphs[0]
    p.clear()
    set_paragraph_rtl(p, align)
    run = p.add_run(text)
    run.bold = bold
    run.font.size = size
    run.font.name = FONT_NAME


def format_effect(delta: int) -> str:
    if delta > 0:
        return f"+{delta}"
    return f"{delta}"


def format_criteria_list(effects: list) -> str:
    if not effects:
        return "—"
    lines = []
    for i, effect in enumerate(effects, 1):
        name = CRITERION_FA.get(effect["criterionKey"], effect["criterionKey"])
        lines.append(f"{i}. {name}")
    return "\n".join(lines)


def format_effects_list(effects: list) -> str:
    if not effects:
        return "—"
    lines = []
    for i, effect in enumerate(effects, 1):
        lines.append(f"{i}. {format_effect(effect['delta'])}")
    return "\n".join(lines)


def set_table_fixed_layout(table):
    tbl = table._tbl
    tbl_pr = tbl.tblPr
    if tbl_pr is None:
        tbl_pr = OxmlElement("w:tblPr")
        tbl.insert(0, tbl_pr)
    layout = OxmlElement("w:tblLayout")
    layout.set(qn("w:type"), "fixed")
    tbl_pr.append(layout)


def main():
    questions = sorted(
        json.loads(QUESTIONS_PATH.read_text(encoding="utf-8")),
        key=lambda q: q["order"],
    )

    doc = Document()
    section = doc.sections[0]
    section.orientation = WD_ORIENT.LANDSCAPE
    section.page_width, section.page_height = section.page_height, section.page_width
    section.left_margin = Cm(1.5)
    section.right_margin = Cm(1.5)
    section.top_margin = Cm(1.5)
    section.bottom_margin = Cm(1.5)

    title = doc.add_paragraph()
    set_paragraph_rtl(title, WD_ALIGN_PARAGRAPH.CENTER)
    title_run = title.add_run("جدول پرسشنامه تطبیقی")
    title_run.bold = True
    title_run.font.size = FONT_SIZE_TITLE
    title_run.font.name = FONT_NAME

    subtitle = doc.add_paragraph()
    set_paragraph_rtl(subtitle, WD_ALIGN_PARAGRAPH.CENTER)
    subtitle_run = subtitle.add_run("قواعد تأثیر پاسخ کاربر بر معیارهای تصمیم‌گیری")
    subtitle_run.font.size = Pt(12)
    subtitle_run.font.name = FONT_NAME

    doc.add_paragraph()

    headers = ["ردیف", "پرسش", "پاسخ کاربر", "معیار تحت تأثیر", "اثر قاعده", "تفسیر"]
    col_widths = [Cm(1.0), Cm(4.5), Cm(3.2), Cm(3.5), Cm(1.8), Cm(6.5)]

    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    set_table_fixed_layout(table)

    for row in table.rows:
        for idx, width in enumerate(col_widths):
            row.cells[idx].width = width

    for idx, header in enumerate(headers):
        cell = table.rows[0].cells[idx]
        write_cell_text(cell, header, bold=True, size=FONT_SIZE_HEADER, align=WD_ALIGN_PARAGRAPH.CENTER)
        shade_cell(cell)

    row_num = 1
    for question in questions:
        qid = question["id"]
        options = question["options"]
        start_row_idx = len(table.rows)

        for opt_idx, option in enumerate(options):
            row = table.add_row()
            effects = option.get("effects", [])
            interpretation = INTERPRETATIONS.get(
                (qid, option["id"]),
                "تغییری در امتیاز تعدیل معیارها ایجاد نمی‌شود.",
            )

            write_cell_text(row.cells[0], str(row_num), align=WD_ALIGN_PARAGRAPH.CENTER)
            write_cell_text(row.cells[2], OPTION_FA.get(option["id"], option["text"]))
            write_cell_text(row.cells[3], format_criteria_list(effects))
            write_cell_text(row.cells[4], format_effects_list(effects), align=WD_ALIGN_PARAGRAPH.CENTER)
            write_cell_text(row.cells[5], interpretation)
            row_num += 1

        first_row = table.rows[start_row_idx]
        question_cell = first_row.cells[1]
        write_cell_text(
            question_cell,
            f"{question['order']}. {QUESTION_FA.get(qid, question['text'])}",
            align=WD_ALIGN_PARAGRAPH.CENTER,
        )

        if len(options) > 1:
            for i in range(1, len(options)):
                question_cell.merge(table.rows[start_row_idx + i].cells[1])

    note = doc.add_paragraph()
    set_paragraph_rtl(note)
    note_run = note.add_run(
        "توضیح: اثر قاعده نشان‌دهنده تغییر امتیاز تعدیل معیار است. "
        "مقدار مثبت اهمیت نسبی معیار را افزایش و مقدار منفی آن را کاهش می‌دهد. "
        "علامت «—» به معنای عدم تأثیر بر معیارها است."
    )
    note_run.font.size = Pt(10)
    note_run.font.name = FONT_NAME
    note_run.italic = True

    OUTPUT_PATH.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUTPUT_PATH)
    print(f"Created: {OUTPUT_PATH}")


if __name__ == "__main__":
    main()
