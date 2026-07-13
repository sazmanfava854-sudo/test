#!/usr/bin/env python3
"""Generate Word document with correct OMML equations for AHP priority vector steps 1-4."""

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import qn
from docx.shared import Pt, Cm
from pathlib import Path

M = "http://schemas.openxmlformats.org/officeDocument/2006/math"
W = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"


def set_rtl(paragraph):
    p_pr = paragraph._p.get_or_add_pPr()
    bidi = OxmlElement("w:bidi")
    bidi.set(qn("w:val"), "1")
    p_pr.append(bidi)


def rtl_para(doc, text, bold=False, size=12):
    p = doc.add_paragraph()
    set_rtl(p)
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    r = p.add_run(text)
    r.bold = bold
    r.font.name = "B Nazanin"
    r.font.size = Pt(size)
    r._element.rPr.rFonts.set(qn("w:ascii"), "B Nazanin")
    r._element.rPr.rFonts.set(qn("w:hAnsi"), "B Nazanin")
    return p


def add_equation(doc, omml_math: str, label: str = ""):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    xml = (
        f'<m:oMathPara xmlns:m="{M}" xmlns:w="{W}">'
        f"<m:oMath>{omml_math}</m:oMath>"
        f"</m:oMathPara>"
    )
    p._p.append(parse_xml(xml))
    if label:
        lp = doc.add_paragraph()
        lp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        lr = lp.add_run(label)
        lr.font.name = "B Nazanin"
        lr.font.size = Pt(11)
        lr.italic = True


def sum_limits(sub_text: str, sup_text: str, body: str) -> str:
    """Summation with visible lower/upper limits."""
    return f"""
    <m:nary>
      <m:naryPr>
        <m:chr m:val="∑"/>
        <m:limLoc m:val="undOvr"/>
        <m:grow m:val="1"/>
      </m:naryPr>
      <m:sub><m:r><m:t>{sub_text}</m:t></m:r></m:sub>
      <m:sup><m:r><m:t>{sup_text}</m:t></m:r></m:sup>
      <m:e>{body}</m:e>
    </m:nary>
    """


# Complete OMML equations (tested structure for Word)
EQ_3_4_1 = """
  <m:sSup>
    <m:e><m:sSub><m:e><m:r><m:t>w</m:t></m:r></m:e><m:sub><m:r><m:t>i</m:t></m:r></m:sub></m:sSub></m:e>
    <m:sup><m:r><m:t>(0)</m:t></m:r></m:sup>
  </m:sSup>
  <m:r><m:t xml:space="preserve"> = </m:t></m:r>
  <m:f>
    <m:num><m:r><m:t>1</m:t></m:r></m:num>
    <m:den><m:r><m:t>n</m:t></m:r></m:den>
  </m:f>
"""

EQ_3_4_2 = f"""
  <m:sSub>
    <m:e><m:r><m:t>w</m:t></m:r></m:e>
    <m:sub><m:r><m:t>i</m:t></m:r></m:sub>
  </m:sSub>
  <m:r><m:t>′</m:t></m:r>
  <m:r><m:t xml:space="preserve"> = </m:t></m:r>
  {sum_limits("j=1", "n", """
    <m:sSub>
      <m:e><m:r><m:t>ā</m:t></m:r></m:e>
      <m:sub><m:r><m:t>ij</m:t></m:r></m:sub>
    </m:sSub>
    <m:r><m:t> · </m:t></m:r>
    <m:sSup>
      <m:e><m:sSub><m:e><m:r><m:t>w</m:t></m:r></m:e><m:sub><m:r><m:t>j</m:t></m:r></m:sub></m:sSub></m:e>
      <m:sup><m:r><m:t>(t)</m:t></m:r></m:sup>
    </m:sSup>
  """)}
"""

EQ_3_4_3 = f"""
  <m:sSup>
    <m:e><m:sSub><m:e><m:r><m:t>w</m:t></m:r></m:e><m:sub><m:r><m:t>i</m:t></m:r></m:sub></m:sSub></m:e>
    <m:sup><m:r><m:t>(t+1)</m:t></m:r></m:sup>
  </m:sSup>
  <m:r><m:t xml:space="preserve"> = </m:t></m:r>
  <m:f>
    <m:num>
      <m:sSub>
        <m:e><m:r><m:t>w</m:t></m:r></m:e>
        <m:sub><m:r><m:t>i</m:t></m:r></m:sub>
      </m:sSub>
      <m:r><m:t>′</m:t></m:r>
    </m:num>
    <m:den>
      {sum_limits("k=1", "n", """
        <m:sSub>
          <m:e><m:r><m:t>w</m:t></m:r></m:e>
          <m:sub><m:r><m:t>k</m:t></m:r></m:sub>
        </m:sSub>
        <m:r><m:t>′</m:t></m:r>
      """)}
    </m:den>
  </m:f>
"""

EQ_3_4_4 = """
  <m:sSub>
    <m:e><m:r><m:t>max</m:t></m:r></m:e>
    <m:sub><m:r><m:t>i</m:t></m:r></m:sub>
  </m:sSub>
  <m:r><m:t xml:space="preserve"> | </m:t></m:r>
  <m:sSup>
    <m:e><m:sSub><m:e><m:r><m:t>w</m:t></m:r></m:e><m:sub><m:r><m:t>i</m:t></m:r></m:sub></m:sSub></m:e>
    <m:sup><m:r><m:t>(t+1)</m:t></m:r></m:sup>
  </m:sSup>
  <m:r><m:t xml:space="preserve"> − </m:t></m:r>
  <m:sSup>
    <m:e><m:sSub><m:e><m:r><m:t>w</m:t></m:r></m:e><m:sub><m:r><m:t>i</m:t></m:r></m:sub></m:sSub></m:e>
    <m:sup><m:r><m:t>(t)</m:t></m:r></m:sup>
  </m:sSup>
  <m:r><m:t xml:space="preserve"> | </m:t></m:r>
  <m:r><m:t> &lt; </m:t></m:r>
  <m:r><m:t>ε</m:t></m:r>
"""

EQ_3_4_5 = f"""
  {sum_limits("i=1", "n", """
    <m:sSub>
      <m:e><m:r><m:t>w</m:t></m:r></m:e>
      <m:sub><m:r><m:t>i</m:t></m:r></m:sub>
    </m:sSub>
  """)}
  <m:r><m:t xml:space="preserve"> = 1</m:t></m:r>
"""


def main():
    out = Path("/workspace/docs/3-4-2_AHP_Priority_Vector_Formulas.docx")
    out.parent.mkdir(parents=True, exist_ok=True)

    doc = Document()
    for s in doc.sections:
        s.page_width, s.page_height = Cm(21), Cm(29.7)
        s.left_margin = s.right_margin = Cm(2.5)

    rtl_para(doc, "۳-۴-۲ محاسبه بردار اولویت", bold=True, size=14)
    doc.add_paragraph()

    rtl_para(
        doc,
        "پس از تشکیل ماتریس تجمیع‌شده Ā، بردار اولویت معیارها با روش تکرار توان "
        "(Power Iteration) محاسبه می‌شود. فرمول‌های مراحل ۱ تا ۴:",
    )
    doc.add_paragraph()

    rtl_para(doc, "۱. مقداردهی اولیه (t = 0)", bold=True)
    add_equation(doc, EQ_3_4_1, "(۳-۴-۱)")

    rtl_para(doc, "۲. ضرب ماتریسی", bold=True)
    add_equation(doc, EQ_3_4_2, "(۳-۴-۲)")

    rtl_para(doc, "۳. نرمال‌سازی", bold=True)
    add_equation(doc, EQ_3_4_3, "(۳-۴-۳)")

    rtl_para(doc, "۴. شرط همگرایی", bold=True)
    add_equation(doc, EQ_3_4_4, "(۳-۴-۴)")
    rtl_para(doc, "که در آن ε = 10⁻¹².")

    rtl_para(doc, "قید نرمال‌سازی بردار نهایی", bold=True)
    add_equation(doc, EQ_3_4_5, "(۳-۴-۵)")

    rtl_para(
        doc,
        "بردار W = [w₁, w₂, …, wₙ]ᵀ تقریب بردار ویژه غالب ماتریس Ā است (n = 8).",
    )

    doc.save(out)
    print(f"Created: {out}")


if __name__ == "__main__":
    main()
